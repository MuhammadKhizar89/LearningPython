# print("Hi");
# a=complex(8,2)
# print (a)
#mutable ->Changes---->List

#immutable ->Doesn't--->tuple

# dictionary->key and value (objects in javascript)

# **->power 2**4 =16  2power4=16
# //->floor division 8//5=1

#Simple Calculator
# by default it takes string

# num1=input("Enter Number:");
# num2=input("Enter Number:");
# print("Addition",int(num1)+int(num2));

# num1=int(input("Enter Number:"));
# num2=int(input("Enter Number:"));
# print("Addition",num1+num2);

# abc="hello" 
# new=abc[0:2] # 2 index sy peechy peechy
# print (new)

# srtipp="!!!hellom!1!2!!"
# print(srtipp.lstrip("!"))
# print(srtipp.rstrip("!"))
# print(srtipp.strip("!"))
# i = "shery"
# for k in range(len(i)):
#     print(i[k])



# # List
# Marks=[2,3,"wq"]
# print (Marks);

# if(Marks[2]=="wq"):
#     print("wq")

# if 2 in Marks:
#     print("Present")

# if "w" in "wq":
#     print("Present");    

# print(Marks[0:4:2]); #staring,ending,skip

# list=[i*i for i in range(10) if i%2==0]
# print(list);

# list.append(5);
# print(list);
# list.sort();
# print(list);
# list.sort(reverse=True);
# print(list);
# list.pop();
# # print(list.index(0));
# # print(list);
# print(list.index(64)); # yani 64 kon sy index pr pra hy
# print(list.count(5));
# m=list;
# m[0]=32;
# print(list)
# # m make changes in list so we use
# m=list.copy();
# m[0]=2;
# print(list)
# list.insert(1,78)#1 index pr 78
# print(list)
# list.extend(m);
# print(list)
# k=list+m;
# print(k);
# print(list)

# # Tuple
# tup=(1,2,"hehe")
# print(tup);
# if 2 in tup:
#     print("Present")

# # slicing krny kay baad new tuple return hoga previous change ni hoga
# tup2=tup[0:2];
# print(tup2);
# # We can cancatinate two tuples
# print(tup.count(3));
# print(tup.index(1,0,2)); #tofind,starting index,ending index
# # to change tuple convert it into list
# # Define a tuple
# GFG_tuple = (1, 2, 3)
 
# # Convert the tuple to a list for changing
# GFG_list = list(GFG_tuple)
# GFG_list.append(2);
# print(GFG_list)
# GFG_tuple=tuple(GFG_list);
# print(GFG_tuple)

# # fString
# a="hello"
# b="world"
# print(f"({a} hehe {b})");
# # or old tradition (a.format(b))
# letter="My Name is{} and I am {} years old"
# age="18"
# name="abc"
# print(letter.format(name,age));
# # format can do more things like floating will be upto 2 decimal

# # DocString
# def square(n):
#     '''Return the square of a number'''
#     return n**2
# print(square(5));

# print(square.__doc__)#It returns the docstring means comment of the function
# # To document our code Only work after defination of function

# #Recursion
# def fab(n):
#     if(n==0):
#         return 0;
#     if(n==1):
#         return 1;
#     return fab(n-1)+fab(n-2);

# for a in range(6):
#  print(fab(a));

# #  Sets
# s={1,2,3,3,"hello","hello"}
# print(s);
# # No gaurantee of order
# ab=set(); #to make empty set if we do this{} then first it thinks it is dictionary
# print(type(s),type(ab))

# for value in s:
#     print(value);
# s2={"hello","World"}
# print(s.union(s2));
# print(s);
# s.update(s2);
# print(s);
# s.intersection_update(s2); #update will the orignal one
# print(s);
# s.symmetric_difference(s2);# for values that are not comn in each others
# s.difference(s2); # A-B vala kaam hy
# s2.isdisjoint(s);
# s2.issubset(s);
# s2.issuperset(s);
# s.add("hehe");
# print(s);
# s.remove("hehe");
# print(s)
# s.discard("hehe"); #remove shows error if that elemet not present but discard not
# print(s)
# a=s.pop();
# print(a)
# del s; #delete full set

# # dictionaries
# dic={
#     "name":"hello",
#     "class":5,
#     "age":18,
#     8:90
#     }
# print(dic["name"]);
# print(dic);
# print(dic.get('hehe')); #it will not  through error if key is not present but simple will through error dic['hehe']
# print(dic.keys());
# print(dic.values());
# print(dic.items());
# dic2={3:5};
# dic.update(dic2);
# print(dic);
# dic.pop(8);
# dic.popitem(); #last iem will be deleted
# dic.clear();
# print(dic);

# #for loop and else
# for i in range(10):
#     print(i)        
# else:
#     print("done")

# # Agr break lga dain loop may tou else ni print hoga

### Seek and tell are like pointers in file that tell moves pointer to start from reading
# with open('file.txt','r') as f:
#     print(f.seek(3))
#     print(f.read())

#truncate to fix bytes in file


# #lambda

# double =lambda x:x*2
# print(double(5))

# # MAP
# l=[1,2,4,6,4,3]
# newl=list(map(lambda x:x*2,l))  
# print(newl)
# # Filter
# def filter_function(x):
#     return x>4;

# newl=list(filter(filter_function,l))
# print(newl)

# # Reduce 
# #it only sum all the values


# # is vs ==
# a=3;
# b=3;
# print(a is b); # exact location of object in memory
# print(a==b);  # value

# # Classes
# class Person:
#     def __init__(self,name,age): #Constructor
#         self.name=name
#         self.age=age

# p1=Person("John",18)
# print(p1.name,p1.age)

# # Decorator

# def greet(fx):
#     def abc():
#         print("hello")
#         fx()
#         print("bye")
#     return abc  

# @greet
# def hello():
#     print("Hi")

# hello();
# # *args for many arguments,**kwargs for dictionary 

# # Getter and setter

# class Person:
#     def __init__(self,name,age): #Constructor
#         self._name=name
#         self._age=age
#     @property
#     def name(self):
#         return self._name
#     @name.setter
#     def name(self,name):
#         self._name=name
#     @property
#     def age(self):
#         return self._age
#     @age.setter
#     def age(self,age):
#         self._age=age
# p1=Person("John",18)
# print(p1.name);# Yani ab hm aik method use kr rhy hn likin lg  rha attribute hy encapsulation
# print(p1.age);
# p1.age=19 #setter


# __abc->private member  to access _classname__abc
# _abc->protected member
# abc->public member


# Static Method
# we can make any method in class
# without sending self  
# we use @staticmethod to make any func


# we use @classmethod to change class atrribute for every object notonly one
# or we can make a alternative constructor using @classmethod


# dir __dict__  help()

# dir
# x=[1,2,3]
# print(dir(x))#To print methods and of list 
# print(x.__add__)

# dict --> if we use object.__dict__ it will print all the attributes as dictionary of that object
# help() it will print all the methods and class of that object

# Super to call parent function
# to cal parent constructor use super().__init__(name,age)

# Magic/Dunder methods
# __str__ method if i write printobject it will print whole object
# i will call str but method name is __str__ so this is magic method

# __call__ method can be call as object() 

# #Exercise 7
# import os
# # os.rename("Images/abc.txt","Images/1.txt")
# k=1
# files=os.listdir("Images")
# for i in files:
#     os.rename(f"Images/{i}",f"Images/{k}.png")
#     k=k+1;

# # Operator Overloading
# #addition  
# v1=Vector(1,2)
# v2=Vector(3,4)
# v1+v2
# # then we make __add__(self,x) x is second v2 and self is first v1 and customize the addition
# # more methods here to overload see documentation



# # Warlus Operator
# # to assign a value in expression
# a=False; 
# if(a:=True):
#     print(a);

# numbers=[1,2,3,4,5];
# while(n:=len(numbers))>0:
#     print(numbers.pop())

# Shutil have may things like to copy one file
# import shutil
# shutil.copy("Learning.py","Learning2.py");

# shutil.copytree->copy whole folder
# shutil.move->move folder
# shutil.rmtree->delete folder

#Exercise 9
# lis=["ayan","jazib","Ali"] 
# import win32com.client as wincom # type: ignore
# import time
# speak = wincom.Dispatch("SAPI.SpVoice")
# for i in lis:
#     text = f"{i} how are you"
#     speak.Speak(text)
# time.sleep(1) 
# text = "byby"
# speak.Speak(text)

# import requests
# response=requests.get("https://www.codewithharry.com")
# print(response.text);

# beautifulsoup for->beautifying the html code or we can find heading and more things that is in
# html code

# # Generator run time generate kry ga yani store ni kray ga memory waste ni ho gi
# def abc():
#  for i in range(5):
#   yield i

# gen=abc()
# or print(next(gen))
# for j in gen:
#  print(j)

# # jb zada kaam ho ga hr value kay liye to hm jis ki zarurat ho gi us ko generate kra lain gy taky memory waste na ho

# Function cashing
# yani agr aik kaam value 2 kay liye function kr rha hy tou ab vo yaad rkhy ga yani dubara vo kaam ni kry ga taky time bach sky
# from functools import lru_cache
# import time
# @lru_cache(maxsize=None)
# def fx(n):
#  time.sleep(2)
#  return n*5;


# print(fx(30))
# print("done for 30")
# print(fx(2))
# print("done for 2")
# print(fx(30))
# print("done for 30")
# print(fx(2))
# print("done for 2")

# # Exercise 10
# import asyncio
# import time
# from desktop_notifier import DesktopNotifier
# notifier = DesktopNotifier()
# async def main():
#     n = await notifier.send(title="Drink Water", message="Sent from Khizar")
#     await asyncio.sleep(5)  # wait a bit before clearing notification
#     await notifier.clear(n)  # removes the notification
#     await notifier.clear_all()  # removes all notifications for this app


# x=time.localtime().tm_hour
# while(True):    
#  if(x!=time.localtime().tm_hour):
#     asyncio.run(main())
#     x=time.localtime().tm_hour;


#Regular Expression simple built in func to find things
# import re
# txt="The rain in Spain"
# x=re.search("ai",txt)
# print(x)  #output-><re.Match object; span=(3, 5), match='ai'>

# Async Io ->like javascript async func and await

# # Multi Threading in Python
# jb hmy rukny ki zarurat ni hy i want to call 10 function at a time 
# jesy async use hota hy kay vo bhi aik sath kaam krny ko khta hy taky baqi kaam hoty rhain

# import threading
# import time

# def print_numbers():
#     for i in range(5):
#         print(i)
#         time.sleep(1)

# def print_letters():
#     for letter in 'abcde':
#         print(letter)
#         time.sleep(1)

# t1 = threading.Thread(target=print_numbers)
# t2 = threading.Thread(target=print_letters)

# t1.start()
# t2.start()

# t1.join()
# t2.join()

# # MultiProcessing same hi lg rha hy
# import time
# import random
# from scholarly import scholarly, ProxyGenerator

# def format_bpp_citation(publication):
#     try:
#         authors = ', '.join(author for author in publication['bib']['author'])
#     except KeyError:
#         authors = "Unknown"
#     try:
#         title = publication['bib']['title']
#     except KeyError:
#         title = "No title available"
#     try:
#         journal = publication['bib'].get('journal', 'No journal available')
#     except KeyError:
#         journal = "No journal available"
#     try:
#         year = publication['bib'].get('pub_year', 'No year available')
#     except KeyError:
#         year = "No year available"
#     try:
#         volume = publication['bib'].get('volume', 'No volume available')
#     except KeyError:
#         volume = "No volume available"
#     try:
#         pages = publication['bib'].get('pages', 'No pages available')
#     except KeyError:
#         pages = "No pages available"
    
#     # Formatting the citation in BPP style
#     return f"{authors} ({year}). {title}. {journal}, {volume}, {pages}."

# def get_first_citation(query, max_retries=5, delay=5):
#     pg = ProxyGenerator()
#     pg.FreeProxies()
#     scholarly.use_proxy(pg)
    
#     for attempt in range(max_retries):
#         try:
#             search_query = scholarly.search_pubs(query)
#             first_result = next(search_query, None)
            
#             if first_result:
#                 # Debug print to inspect the structure
#                 # print(first_result)
#                 citation = format_bpp_citation(first_result)
#                 return citation
#             else:
#                 return "No results found."
#         except Exception as e:
#             print(f"Attempt {attempt + 1} failed: {e}")
#             time.sleep(delay + random.uniform(0, 3))  # Adding some randomness to delay

#     return "Failed to retrieve results after multiple attempts."

# # Example usage
# query = "Mars is an international manufacturer and marketer of confection and food products with some of its popular brands including M&M’s, Snickers, and Twix"
# citation = get_first_citation(query)
# print(citation)
from docx import Document
import nltk
import random
import time
from scholarly import scholarly, ProxyGenerator

nltk.download('punkt')  # Ensure NLTK tokenizer is downloaded for sentence tokenization

def extract_sentences(doc_path):
    doc = Document(doc_path)
    sentences = []
    for para in doc.paragraphs:
        if para.text.strip():  # Ensure paragraph is not empty
            para_sentences = nltk.sent_tokenize(para.text.strip())
            if len(para_sentences) > 2:  # For long paragraphs
                first_sentence = para_sentences[0]
                middle_sentence = para_sentences[len(para_sentences) // 2]
                last_sentence = para_sentences[-1]
                sentences.append((first_sentence, middle_sentence, last_sentence))
            else:  # For short paragraphs
                first_sentence = para_sentences[0]
                sentences.append((first_sentence,))
    return sentences

def format_bpp_citation(publication):
    try:
        authors = ', '.join(author for author in publication['bib']['author'])
    except KeyError:
        authors = "Unknown"
    try:
        title = publication['bib']['title']
    except KeyError:
        title = "No title available"
    try:
        journal = publication['bib'].get('journal', 'No journal available')
    except KeyError:
        journal = "No journal available"
    try:
        year = publication['bib'].get('pub_year', 'No year available')
    except KeyError:
        year = "No year available"
    try:
        volume = publication['bib'].get('volume', 'No volume available')
    except KeyError:
        volume = "No volume available"
    try:
        pages = publication['bib'].get('pages', 'No pages available')
    except KeyError:
        pages = "No pages available"
    # Formatting the citation in BPP style
    print(f"{authors} ({year}). {title}. {journal}, {volume}, {pages}.")    
    return f"{authors} ({year}). {title}. {journal}, {volume}, {pages}."

def get_first_citation(query, max_retries=5, delay=5):
    pg = ProxyGenerator()
    pg.FreeProxies()
    scholarly.use_proxy(pg)
    
    for attempt in range(max_retries):
        try:
            print(query)
            search_query = scholarly.search_pubs(query)
            first_result = next(search_query, None)
            if first_result:
                return first_result  # Return the publication object
            else:
                return None  # Return None if no results found
        except Exception as e:
            print(f"Attempt {attempt + 1} failed: {e}")
            time.sleep(delay + random.uniform(0, 3))  # Adding some randomness to delay

    return None  # Return None if failed to retrieve results after multiple attempts

# Function to update paragraphs with citations
def update_docx_with_citations(doc_path, sentences, citations):
    doc = Document(doc_path)
    
    for para, sentence_group, citation_group in zip(doc.paragraphs, sentences, citations):
        if para.text.strip():  # Ensure paragraph is not empty
            para_text = para.text.strip()
            para_sentences = nltk.sent_tokenize(para_text)
            new_para_text = para_sentences.copy()
            for i, sentence in enumerate(sentence_group):
                if sentence in para_sentences:
                    sentence_index = para_sentences.index(sentence)
                    new_para_text[sentence_index] += f". {citation_group[i]}"
            para.clear()  # Clear existing content
            for sentence in new_para_text:
                para.add_run(sentence + " ")
    doc.save(doc_path)

# Example usage
doc_path = 'path_to_your_document.docx'  # Replace with your actual document path
sentences = extract_sentences(doc_path)

citations = []
for sentence_group in sentences:
    group_citations = []
    for sentence in sentence_group:
        query = sentence
        publication = get_first_citation(query)
        if publication:
            citation = format_bpp_citation(publication)
            group_citations.append(citation)  # Store formatted citation
        else:
            group_citations.append("")  # Handle cases where no citation is found
    citations.append(group_citations)

# Update DOCX file with citations
update_docx_with_citations(doc_path, sentences, citations);
print(f"Citations have been inserted into '{doc_path}' successfully.");
